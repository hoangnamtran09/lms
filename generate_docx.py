#!/usr/bin/env python3
"""Convert KICH_BAN_THUYET_TRINH.md to Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Read markdown
with open('/Users/hoangnamtran/WORKS/lms/KICH_BAN_THUYET_TRINH.md', 'r') as f:
    text = f.read()

lines = text.split('\n')

i = 0
while i < len(lines):
    line = lines[i]

    # Skip empty lines
    if not line.strip():
        i += 1
        continue

    # H1 - Title
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        i += 1
        continue

    # H2 - Section headers (slide titles)
    if line.startswith('## SLIDE') or line.startswith('## Phụ lục'):
        title = line[3:].strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        i += 1
        continue

    # Separator
    if line.strip() == '---':
        p = doc.add_paragraph()
        run = p.add_run('—' * 40)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    # Bold text paragraphs (stage directions: [▶], (tạm dừng), *italics*)
    is_meta = line.strip().startswith('>') or line.strip().startswith('*') or '[▶]' in line or '(tạm dừng)' in line

    # Process inline formatting
    p = doc.add_paragraph()
    if is_meta:
        p.paragraph_format.left_indent = Cm(1)

    # Parse bold **text** and italic *text*
    remaining = line
    while remaining:
        # Bold
        bold_match = re.match(r'^(.*?)\*\*(.+?)\*\*(.*)', remaining, re.DOTALL)
        if bold_match:
            if bold_match.group(1):
                run = p.add_run(bold_match.group(1))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
            run = p.add_run(bold_match.group(2))
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            # Color for slide transitions
            if '[▶]' in bold_match.group(2):
                run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
            remaining = bold_match.group(3)
            continue

        # Inline code or italic markdown (e.g., *(tạm dừng)*)
        italic_match = re.match(r'^(.*?)\*(.+?)\*(.*)', remaining, re.DOTALL)
        if italic_match:
            if italic_match.group(1):
                run = p.add_run(italic_match.group(1))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
            run = p.add_run(italic_match.group(2))
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            remaining = italic_match.group(3)
            continue

        # Normal text
        run = p.add_run(remaining)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        break

    i += 1

# Save
output = '/Users/hoangnamtran/WORKS/lms/KICH_BAN_THUYET_TRINH.docx'
doc.save(output)
print(f'Saved to {output}')
